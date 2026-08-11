from docx import Document
from docx.shared import Inches
from datetime import datetime
import os

def build_docx(document_type: str, sections: dict, chart_paths: list = None, output_dir: str = "outputs") -> str:
    """
    Takes generated section content and assembles a real Word document.
    Optionally embeds chart images (e.g. budget/timeline PNGs) after the sections.
    Returns the file path of the saved .docx.
    """
    doc = Document()

    # Title
    title = document_type.replace('_', ' ').title()
    doc.add_heading(title, level=0)

    # Each section becomes a heading + paragraph
    for step_name, content in sections.items():
        doc.add_heading(step_name, level=1)
        doc.add_paragraph(content)

    # Charts, if any were generated successfully
    if chart_paths:
        doc.add_heading("Visual Summary", level=1)
        for chart_path in chart_paths:
            if os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(5.5))

    # Save with a unique filename
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{document_type}_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    return filepath